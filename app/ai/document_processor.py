"""
Simplified document processor for PDF, DOCX, and TXT files.
"""

import os
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.core.config import get_settings
from app.core.logging import get_logger
from app.ai import vector_store

logger = get_logger("ai.document_processor")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="cp1251") as f:
            return f.read()


def extract_text(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    ext = Path(file_path).suffix.lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".txt", ".md"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def split_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """Split text into chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    return splitter.split_text(text)


async def process_document(
    file_path: str,
    source_name: Optional[str] = None
) -> dict:
    """Process a document and add it to the vector store."""
    try:
        # Extract text
        text = extract_text(file_path)
        
        if not text.strip():
            return {
                "success": False,
                "error": "Document is empty or could not be read"
            }
        
        # Split into chunks
        chunks = split_text(text)
        
        # Create documents
        source = source_name or Path(file_path).name
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_index": i
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # Add to vector store
        ids = await vector_store.add_documents(
            documents,
            metadata={"source": source}
        )
        
        logger.info(f"Processed document {source}: {len(chunks)} chunks")
        
        return {
            "success": True,
            "source": source,
            "chunks_count": len(chunks),
            "vector_ids": ids
        }
        
    except Exception as e:
        logger.error(f"Failed to process document: {e}")
        return {
            "success": False,
            "error": str(e)
        }


async def delete_document(source_name: str) -> dict:
    """Delete a document from the vector store."""
    try:
        success = await vector_store.delete_by_source(source_name)
        
        return {
            "success": success,
            "source": source_name
        }
        
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        return {
            "success": False,
            "error": str(e)
        }


def save_uploaded_file(file_content: bytes, filename: str) -> str:
    """Save uploaded file to disk."""
    settings = get_settings()
    
    # Create upload directory if it doesn't exist
    upload_dir = Path(settings.upload_dir)
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    # Save file
    file_path = upload_dir / filename
    with open(file_path, "wb") as f:
        f.write(file_content)
    
    return str(file_path)


def get_allowed_extensions() -> List[str]:
    """Get list of allowed file extensions."""
    return [".pdf", ".docx", ".txt", ".md"]
